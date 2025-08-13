# minNode.py - Simplified OpenAI Integration (with better diagnostics)
import os
import json
from typing import Dict, List, Any, Tuple, Optional
from openai import OpenAI
from PyPDF2 import PdfReader
import docx

class OpenAIClient:
    def __init__(self, api_key: Optional[str] = None, base_url: Optional[str] = None):
        self.client: Optional[OpenAI] = None
        # Allow override via env var if present
        self.model = os.getenv("OPENAI_MODEL", "gpt-4o-mini")
        self.last_error: Optional[str] = None
        # Initialize client immediately if a key exists in args or env
        key = api_key or os.getenv("OPENAI_API_KEY")
        self._init_client(key, base_url or os.getenv("OPENAI_BASE_URL"))

    def _init_client(self, api_key: Optional[str], base_url: Optional[str] = None) -> None:
        try:
            if not api_key:
                self.client = None
                self.last_error = "OPENAI_API_KEY is missing"
                return
            if base_url:
                self.client = OpenAI(api_key=api_key, base_url=base_url)
            else:
                self.client = OpenAI(api_key=api_key)
            self.last_error = None
        except Exception as e:
            self.client = None
            self.last_error = f"Failed to initialize OpenAI client: {e}"

    def set_api_key(self, api_key: str) -> bool:
        self._init_client(api_key)
        return self.client is not None

    def test_connection(self) -> Tuple[bool, str]:
        """Return (ok, message). On failure, message contains the exception."""
        if not self.client:
            # Ensure a tuple is returned even if the client is not initialized
            return False, (self.last_error or "OpenAI client is not initialized")
        try:
            # On success, return True and a success message
            self.client.models.list()
            return True, "Connection successful"
        except Exception as e:
            self.last_error = str(e)
            # On exception, return False and the error message
            return False, self.last_error

    def generate_meeting_minutes(self, transcript: str) -> Dict[str, Any]:
        if not self.client:
            # Return fallback, but preserve the reason
            return self._create_fallback_minutes()

        try:
            response = self.client.chat.completions.create(
                model=self.model,
                messages=[
                    {"role": "system", "content": self._get_system_prompt()},
                    {"role": "user", "content": f"Analyze this meeting transcript:\n\n{transcript}"}
                ],
                temperature=0.2,
                max_tokens=3000,
                response_format={"type": "json_object"}
            )
            minutes = json.loads(response.choices[0].message.content)
            return self._clean_minutes(minutes)
        except Exception as e:
            self.last_error = str(e)
            return self._create_fallback_minutes()

    def _get_system_prompt(self) -> str:
        # This new prompt is more flexible and better suited for various conversation types.
        return """You are an intelligent assistant that analyzes conversations and generates a structured summary in JSON format. Your goal is to extract the most important information, focusing on clarity and factual accuracy based on the provided transcript.

Required JSON structure:
{
  "summary": "A concise, 2-4 sentence summary of the entire conversation's purpose and flow.",
  "participants": ["Name 1 (Role, if specified)", "Name 2 (Role, if specified)"],
  "discussion_points": [
    "A key topic or question that was discussed.",
    "Another significant point of discussion."
  ],
  "outcomes_or_decisions": [
    "Any final decisions, conclusions, or results from the conversation."
  ],
  "next_steps": [
    "Any explicit mentions of future actions or follow-ups."
  ]
}

Rules:
- ALWAYS output a valid JSON object.
- If a section has no relevant information (e.g., no decisions were made), use an empty list [].
- Extract participant names and their roles if mentioned (e.g., "Cate (Material Science)").
- Keep summaries and points concise and directly from the transcript.
- Do not invent or infer information not present in the text. Focus on what was actually said."""


    def _clean_minutes(self, minutes: Dict) -> Dict[str, Any]:
        # Updated to match the new, more flexible JSON structure from the prompt.
        return {
            "summary": str(minutes.get("summary", "")).strip() or "No summary available.",
            "participants": [str(p).strip() for p in minutes.get("participants", []) if str(p).strip()],
            "discussion_points": [str(t).strip() for t in minutes.get("discussion_points", []) if str(t).strip()],
            "outcomes_or_decisions": [str(d).strip() for d in minutes.get("outcomes_or_decisions", []) if str(d).strip()],
            "next_steps": [str(s).strip() for s in minutes.get("next_steps", []) if str(s).strip()]
        }

    # The _clean_action_items method is no longer needed with this new structure and can be removed.
    def _clean_action_items(self, items: List) -> List[Dict]:
        return [] # Returning empty list for compatibility, but it's unused.

    def _create_fallback_minutes(self) -> Dict[str, Any]:
        return {
            "summary": "Could not automatically generate minutes. Manual review required.",
            "participants": [],
            "key_topics": ["Manual review required"],
            "decisions": [],
            "action_items": [],
            "next_steps": ["Review transcript manually"]
        }

# File processing functions
def extract_text_from_pdf(file) -> str:
    try:
        reader = PdfReader(file)
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text and text.strip():
                text_parts.append(text.strip())
        return "\n\n".join(text_parts)
    except Exception:
        return ""

def extract_text_from_docx(file) -> str:
    try:
        doc = docx.Document(file)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        return "\n".join(text_parts)
    except Exception:
        return ""